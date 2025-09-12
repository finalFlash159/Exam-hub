import os
import logging
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def extract_text(self, file_path):
        """Extract text content from PDF or DOCX files"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF files"""
        extracted_text = ""
        
        try:
            logger.info(f"Processing PDF file: {file_path}")
            # Open the PDF file
            with fitz.open(file_path) as pdf_document:
                # Iterate through each page
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    page_text = page.get_text()
                    extracted_text += page_text
                    
            logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX files"""
        extracted_text = ""
        
        try:
            logger.info(f"Processing DOCX file: {file_path}")
            # Open the DOCX file
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " | "
                    extracted_text += "\n"
                    
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
