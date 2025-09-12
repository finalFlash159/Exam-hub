import os
import logging

try:
    import docx
except ImportError:
    docx = None

from .base import BaseProcessor, ProcessingResult

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseProcessor):
    """Simple DOCX text extractor"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.docx'}
        
        if docx is None:
            raise ImportError("python-docx required: pip install python-docx")
    
    async def extract_text(self, file_path: str) -> ProcessingResult:
        """Extract text from DOCX - simple and clean"""
        try:
            if not os.path.exists(file_path):
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message=f"File not found: {file_path}"
                )
            
            doc = docx.Document(file_path)
            text = ""
            
            # Get text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            text = text.strip()
            if not text:
                return ProcessingResult(
                    success=False,
                    content="",
                    error_message="No text content found in DOCX"
                )
            
            logger.info(f"Extracted {len(text)} chars from DOCX")
            return ProcessingResult(success=True, content=text)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ProcessingResult(
                success=False,
                content="",
                error_message=f"DOCX processing error: {str(e)}"
            )